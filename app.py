from flask import Flask, request, jsonify
from flask_sqlalchemy import SQLAlchemy
from docx import Document

app = Flask(__name__)

# Настройка базы данных
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///testcases.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)


# Модели базы данных
class TestCase(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.String(500), nullable=False)
    precondition = db.Column(db.String(500), nullable=True)  # Добавлено поле для предусловия
    postcondition = db.Column(db.String(500), nullable=True)  # Добавлено поле для постусловия
    comment = db.Column(db.String(500), nullable=True)  # Добавлено поле для комментария


class Step(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    test_case_id = db.Column(db.Integer, db.ForeignKey('test_case.id'), nullable=False)
    description = db.Column(db.String(500), nullable=False)
    expected_result = db.Column(db.String(500), nullable=False)


# Создание базы данных (если базы данных нет)
with app.app_context():
    db.create_all()


# API для работы с тест-кейсами
@app.route('/')
def index():
    return "TMS API is running!"  # Главная страница


@app.route('/testcases', methods=['POST'])
def create_test_case():
    data = request.json
    test_case = TestCase(
        name=data['name'],
        description=data['description'],
        precondition=data.get('precondition'),  # Предусловие
        postcondition=data.get('postcondition'),  # Постусловие
        comment=data.get('comment')  # Комментарий
    )
    db.session.add(test_case)
    db.session.commit()
    return jsonify({"message": "Test case created!", "id": test_case.id}), 201


@app.route('/steps', methods=['POST'])
def add_step():
    data = request.json
    step = Step(
        test_case_id=data['test_case_id'],
        description=data['description'],
        expected_result=data['expected_result']
    )
    db.session.add(step)
    db.session.commit()
    return jsonify({"message": "Step added!"}), 201


@app.route('/export/<int:test_case_id>', methods=['GET'])
def export_to_word(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    steps = Step.query.filter_by(test_case_id=test_case_id).all()

    # Генерация документа Word
    doc = Document()
    doc.add_heading(test_case.name, level=1)
    doc.add_paragraph(test_case.description)

    # Добавляем предусловие, постусловие и комментарий
    if test_case.precondition:
        doc.add_paragraph(f"Предусловие: {test_case.precondition}")
    if test_case.postcondition:
        doc.add_paragraph(f"Постусловие: {test_case.postcondition}")
    if test_case.comment:
        doc.add_paragraph(f"Комментарий: {test_case.comment}")

    # Добавляем таблицу с шагами
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Шаги'
    hdr_cells[1].text = 'Ожидаемый результат'
    hdr_cells[2].text = 'Полученный результат'

    for step in steps:
        row_cells = table.add_row().cells
        row_cells[0].text = step.description
        row_cells[1].text = step.expected_result
        row_cells[2].text = ''  # Оставляем пустым

    filename = f"TestCase_{test_case_id}.docx"
    doc.save(filename)
    return jsonify({"message": f"Exported to {filename}"}), 200


if __name__ == '__main__':
    app.run(debug=True)
